import streamlit as st
import pandas as pd
from docx import Document
import openai
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import requests

# Initialize OpenAI API key
openai.api_key = "sk-proj-vQk9SyZY9GOSxm09KWxtT3BlbkFJBhTI29nx4Y13rZZ3GbDs"

# Function to parse and print tables from the document
def parse_and_print_tables(doc):
    for idx, table in enumerate(doc.tables, start=1):
        st.write(f"Table {idx}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            st.write(row_data)

# Function to format cell value based on the column index and table index
def format_cell_value(value, col_index, table_index=None):
    if pd.isna(value) or value == '':
        return ''
    if table_index == 2 and col_index == 3:
        return str(value)
    if col_index == 0:
        return pd.to_datetime(value).strftime('%b-%y') if not isinstance(value, str) else value
    elif col_index == 2:
        try:
            return f"{int(value)}"
        except ValueError:
            return value
    elif col_index == 3:
        return f"${float(value):,.2f}" if isinstance(value, (int, float)) else value
    elif col_index == 4 or col_index == 5:
        return f"${float(value):,.2f}" if isinstance(value, (int, float)) else value
    elif col_index == 6:
        return f"{int(value)}%" if isinstance(value, (int, float)) else value
    else:
        return str(value)

# Function to update table with new values
def update_table_with_values(table, new_values, start_row=1):
    num_rows_to_add = len(new_values) - (len(table.rows) - start_row)
    if num_rows_to_add > 0:
        for _ in range(num_rows_to_add):
            table.add_row()
    for i, row in enumerate(new_values):
        for j, cell_value in enumerate(row):
            if (i + start_row) < len(table.rows) and j < len(table.columns):
                cell = table.cell(i + start_row, j)
                formatted_value = format_cell_value(cell_value, j)
                cell.text = formatted_value
                cell.width = cell.width
                cell.paragraphs[0].alignment = cell.paragraphs[0].alignment
            else:
                st.write(f"Skipping value at row {i + start_row}, column {j} - out of bounds")

# Function to update specific column of the table with new values
def update_table_column_with_values(table, column_index, new_values, start_row=1, table_index=None):
    num_rows_to_add = len(new_values) - (len(table.rows) - start_row)
    if num_rows_to_add > 0:
        for _ in range(num_rows_to_add):
            table.add_row()
    for i, value in enumerate(new_values):
        row_index = i + start_row
        if row_index < len(table.rows):
            cell = table.cell(row_index, column_index)
            formatted_value = format_cell_value(value, column_index, table_index)
            cell.text = formatted_value
            cell.width = cell.width
            cell.paragraphs[0].alignment = cell.paragraphs[0].alignment
        else:
            st.write(f"Skipping value at row {row_index}, column {column_index} - out of bounds")

    if len(new_values) > 1 and len(table.rows) == start_row + 1:
        for i, value in enumerate(new_values[1:], start=start_row + 1):
            cell = table.cell(i, column_index)
            formatted_value = format_cell_value(value, column_index, table_index)
            cell.text = formatted_value
            cell.width = cell.width
            cell.paragraphs[0].alignment = cell.paragraphs[0].alignment

# Function to set cell border
def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f'w:{edge}'
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in edge_data:
                element.set(qn(key), str(edge_data[key]))

# Main function to create the Streamlit app
def main():
    st.set_page_config(page_title="Monthly Report Updater", page_icon=":bar_chart:", layout="wide")

    st.markdown(
        """
        <style>
        .main {
            background-color: #121212;
        }
        .report-uploader {
            background-color: #ffffff;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1);
        }
        .header {
            text-align: center;
            font-size: 2em;
            font-weight: bold;
            margin-bottom: 30px;
        }
        .subheader {
            font-size: 1.2em;
            font-weight: bold;
            margin-bottom: 10px;
        }
        .instructions {
            background-color: #4CAF50;
            font-size: 1em;
            margin-bottom: 20px;
        }
        .button {
            background-color: #4CAF50;
            color: white;
            border: none;
            padding: 10px 20px;
            text-align: center;
            font-size: 1em;
            border-radius: 5px;
            cursor: pointer;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.markdown("<div class='header'>Monthly Report Generator</div>", unsafe_allow_html=True)

    with st.container():
        col1, col2, col3 = st.columns([1, 2, 1])

        with col2:
            st.markdown("<div class='report-uploader'>", unsafe_allow_html=True)
            st.markdown("<div class='subheader'>Upload Files</div>", unsafe_allow_html=True)
            st.markdown("<div class='instructions'>Please upload the required files in the specified order.</div>", unsafe_allow_html=True)

            timesheet_file = st.file_uploader("1. Upload Sample Timesheet Report.xlsx", type="xlsx", key="timesheet")
            funding_file = st.file_uploader("2. Upload Sample Funding Status.xlsx", type="xlsx", key="funding")
            
            # Removed the file uploader for SAMPLE REPORT.docx and replaced it with loading from URL
            url = "https://github.com/nrjunejo/report-streamlit-app/blob/main/SAMPLE%20REPORT.docx?raw=true"
            response = requests.get(url)
            report_file = BytesIO(response.content)

            month = st.text_input("Enter the month:", key="month")
            year = st.text_input("Enter the year:", key="year")

            if st.button("Update Report"):
                if timesheet_file and funding_file and report_file and month and year:
                    excel_sheet1 = pd.read_excel(timesheet_file)
                    excel_sheet2 = pd.read_excel(funding_file)
                    doc = Document(report_file)

                    new_month = month
                    new_year = year
                    new_date = "1st"

                    for paragraph in doc.paragraphs:
                        if "For the month of" in paragraph.text and "Precision commenced services on" in paragraph.text:
                            updated_sentence = f"For the month of {new_month} {new_year} – Precision commenced services on {new_month} {new_date} {new_year}. Primary PAS attendant continued government security clearance process."
                            paragraph.text = updated_sentence

                    table2_column4_values = excel_sheet1.iloc[26:30, 1].tolist()
                    table = doc.tables[1]
                    update_table_column_with_values(table, 3, table2_column4_values, table_index=2)

                    pas_hours_value = excel_sheet1.iloc[26, 1]
                    pas_backup_hours_value = excel_sheet1.iloc[27, 1]
                    total_hours_value = excel_sheet1.iloc[26:28, 1].sum()

                    table3 = doc.tables[2]
                    num_shifts = sum(1 for row in table3.rows[1:] if row.cells[4].text.strip() == '2')

                    updated_column2_text = f"PAS Assistance Scheduled and provided in {new_month} {new_year} for employee work location at White House Campus.\nPAS Hours = {pas_hours_value}\nPAS Backup Hours = {pas_backup_hours_value}\nTotal Hours= {total_hours_value}\n*Backup = first 2 Hours per {num_shifts} shift"
                    table = doc.tables[0]
                    update_table_column_with_values(table, 1, updated_column2_text.split('\n'))

                    date_range_text = f"{new_month} 1st {new_year} thru {new_month} 30th {new_year}\n{pas_hours_value} (PAS) + {pas_backup_hours_value} (Backup)=\n{total_hours_value} Hours"
                    date_range_lines = date_range_text.split('\n')
                    update_table_column_with_values(table, 3, date_range_lines, start_row=1)

                    table4_values = excel_sheet1.iloc[26:31, 1].values.tolist()
                    table = doc.tables[3]
                    update_table_column_with_values(table, 1, table4_values, start_row=1)

                    table5_values = excel_sheet2.iloc[3:, :7].dropna(how='all').values.tolist()
                    table = doc.tables[4]
                    num_rows_to_add = len(table5_values) - (len(table.rows) - 1)
                    if num_rows_to_add > 0:
                        for _ in range(num_rows_to_add):
                            table.add_row()

                    for i, row_values in enumerate(table5_values):
                        for j, value in enumerate(row_values):
                            cell = table.cell(i+1, j)
                            if j == 6:
                                formatted_value = f"{value * 100:.0f}%" if pd.notna(value) else ''
                            else:
                                formatted_value = format_cell_value(value, j)
                            cell.text = formatted_value
                            cell.width = cell.width
                            cell.paragraphs[0].alignment = cell.paragraphs[0].alignment

                    for j, header in enumerate(table.rows[0].cells):
                        for paragraph in header.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    table = doc.tables[2]
                    for row in table.rows[:]:
                        table._element.remove(row._element)
                    num_rows_to_add = 27
                    for _ in range(num_rows_to_add):
                        table.add_row()

                    headers = ['DAY OF SERVICE', 'DATE OF SERVICE PROVIDED', 'SHIFT START', 'SHIFT END', 'HOURS']
                    for j, header in enumerate(headers):
                        cell = table.cell(0, j)
                        cell.text = header
                        set_cell_border(cell, top={"w:val": "single", "w:sz": "4"}, bottom={"w:val": "single", "w:sz": "4"},
                                        left={"w:val": "single", "w:sz": "4"}, right={"w:val": "single", "w:sz": "4"})
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    table3_values = excel_sheet1.iloc[0:25, :5].values.tolist()
                    for i, row_values in enumerate(table3_values):
                        for j, value in enumerate(row_values):
                            if (i + 1) < len(table.rows) and j < len(table.columns):
                                cell = table.cell(i + 1, j)
                                if j == 1:
                                    try:
                                        formatted_value = pd.to_datetime(value).strftime('%m/%d/%Y') if not pd.isna(value) else ''
                                    except (ValueError, TypeError):
                                        formatted_value = str(value)
                                elif j == 4:
                                    try:
                                        formatted_value = f"{int(value):,d}"
                                    except ValueError:
                                        formatted_value = str(value)
                                else:
                                    formatted_value = format_cell_value(value, j)
                                cell.text = formatted_value
                                set_cell_border(cell, top={"w:val": "single", "w:sz": "4"}, bottom={"w:val": "single", "w:sz": "4"},
                                                left={"w:val": "single", "w:sz": "4"}, right={"w:val": "single", "w:sz": "4"})
                                cell.width = cell.width
                                cell.paragraphs[0].alignment = cell.paragraphs[0].alignment
                            else:
                                st.write(f"Skipping value at row {i + 1}, column {j} - out of bounds")

                    doc.save("Updated_SAMPLE_REPORT.docx")
                    st.success("Report updated successfully!")

                    with open("Updated_SAMPLE_REPORT.docx", "rb") as f:
                        st.download_button("Download Updated Report", f, file_name="Updated_SAMPLE_REPORT.docx")
                else:
                    st.error("Please upload all required files and enter the month and year.")

            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown(
        """
        <style>
        footer {visibility: hidden;}
        .css-1aumxhk {padding: 20px;}
        .css-12oz5g7 {padding-top: 3.5rem;}
        </style>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
